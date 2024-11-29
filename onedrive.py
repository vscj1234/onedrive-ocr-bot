import msal
from flask import Flask, request, redirect, session, url_for, jsonify, render_template
import requests
import os
from io import BytesIO
from docx import Document
import chromadb
from typing import List, Tuple
from openai import OpenAI
import zipfile
import json
import hashlib
from dotenv import load_dotenv
import mimetypes
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
import logging
from ocr import OCRProcessor

logging.basicConfig(level=logging.DEBUG)

load_dotenv()
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=1)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Azure AD configuration
CLIENT_ID = os.environ.get("CLIENT_ID")
CLIENT_SECRET = os.environ.get("CLIENT_SECRET")
TENANT_ID = os.environ.get("TENANT_ID")
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
REDIRECT_URI = 'http://localhost:5000/getAToken'
SCOPE = ['User.Read', 'Files.Read', 'Files.ReadWrite']
# OpenAI configuration
client = OpenAI(api_key=OPENAI_API_KEY)
msal_app = msal.ConfidentialClientApplication(CLIENT_ID, authority=AUTHORITY, client_credential=CLIENT_SECRET)

# Initialize Chroma
chroma_client = chromadb.Client()
collection = chroma_client.create_collection(name="onedrive_docs")

# Initialize Langchain components
embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
vectorstore = Chroma(embedding_function=embeddings, persist_directory="./chromaonedrive_db")

# Initialize OCR processor
ocr_processor = OCRProcessor()

def read_docx(file_path_or_bytes) -> str:
    """
    Read content from a .docx file
    
    Args:
        file_path_or_bytes: Either a file path or bytes/BytesIO object containing the document
        
    Returns:
        str: Extracted text content from the document
    """
    try:
        file_stream = BytesIO(file_path_or_bytes) if isinstance(file_path_or_bytes, bytes) else file_path_or_bytes
        file_stream.seek(0)
        doc = Document(file_stream)
        full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        full_text += [cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
        return " ".join(full_text)
    except (ValueError, zipfile.BadZipFile) as e:
        raise ValueError(f"Invalid document format: {str(e)}")

def fetch_file_content(file_id: str, access_token: str) -> Tuple[str, bytes, str]:
    headers = {'Authorization': f'Bearer {access_token}'}
    file_data = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}', headers=headers).json()
    content_response = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/content', headers=headers)
    return file_data['name'], content_response.content, file_data.get('file', {}).get('mimeType', '')

def get_file_metadata(file_id, access_token):
    headers = {'Authorization': f'Bearer {access_token}'}
    file_data = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}', headers=headers).json()
    return {'id': file_id, 'name': file_data['name'], 'lastModifiedDateTime': file_data['lastModifiedDateTime']}

def load_processed_files():
    if os.path.exists('processed_files.json'):
        with open('processed_files.json', 'r') as f:
            return json.load(f)
    return {}

def save_processed_files(processed_files):
    with open('processed_files.json', 'w') as f:
        json.dump(processed_files, f)

class FileProcessingCache:
    def __init__(self, cache_dir="./cache"):
        self.cache_dir = cache_dir
        self.cache_file = os.path.join(cache_dir, "file_cache.json")
        os.makedirs(self.cache_dir, exist_ok=True)
        self.cache = self._load_cache()

    def _load_cache(self):
        if os.path.exists(self.cache_file):
            with open(self.cache_file, 'r') as f:
                return json.load(f)
        return {}

    def _save_cache(self):
        with open(self.cache_file, 'w') as f:
            json.dump(self.cache, f)

    def _calculate_file_hash(self, file_content: bytes) -> str:
        return hashlib.md5(file_content).hexdigest()

    def needs_processing(self, file_id: str, file_content: bytes, last_modified: str) -> bool:
        file_hash = self._calculate_file_hash(file_content)
        cached_info = self.cache.get(file_id, {})
        return (file_id not in self.cache or 
                cached_info['hash'] != file_hash or 
                cached_info['last_modified'] != last_modified)

    def update_cache(self, file_id: str, file_content: bytes, metadata: dict):
        self.cache[file_id] = {
            'hash': self._calculate_file_hash(file_content),
            'last_modified': metadata['lastModifiedDateTime'],
            'processed_date': datetime.now().isoformat(),
            'metadata': metadata
        }
        self._save_cache()

def process_files_optimized(access_token):
    cache = FileProcessingCache()
    headers = {'Authorization': f'Bearer {access_token}'}
    processed_count, skipped_count = 0, 0

    def process_item_optimized(item, parent_path=""):
        nonlocal processed_count, skipped_count
        item_id = item['id']
        item_metadata = get_file_metadata(item_id, access_token)
        item_path = os.path.join(parent_path, item_metadata['name'])

        if 'folder' in item:
            folder_content = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{item_id}/children', headers=headers).json()
            for child_item in folder_content.get('value', []):
                process_item_optimized(child_item, item_path)
        else:
            try:
                file_name, file_content, file_type = fetch_file_content(item_id, access_token)
                
                # Skip if file hasn't changed
                if not cache.needs_processing(item_id, file_content, item_metadata['lastModifiedDateTime']):
                    skipped_count += 1
                    return

                text_content = None
                if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    text_content = read_docx(file_content)
                elif file_type.startswith('image/') and ocr_processor.is_image_file(file_name):
                    text_content = ocr_processor.extract_text(file_content)

                if text_content:
                    chunks = text_splitter.split_text(text_content)
                    metadatas = [{
                        "source": item_path,
                        "file_name": file_name,
                        "file_id": item_id,
                        "file_type": file_type,
                        "chunk_index": i,
                        "processing_date": datetime.now().isoformat()
                    } for i in range(len(chunks))]
                    
                    vectorstore.add_texts(texts=chunks, metadatas=metadatas, ids=[f"{item_id}_chunk_{i}" for i in range(len(chunks))])
                    cache.update_cache(item_id, file_content, item_metadata)
                    processed_count += 1

            except Exception as e:
                raise

    root_items = requests.get('https://graph.microsoft.com/v1.0/me/drive/root/children', headers=headers).json()
    for item in root_items.get('value', []):
        process_item_optimized(item)
    return processed_count, skipped_count

# Initialize ConversationalRetrievalChain
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, temperature=0.7, max_tokens=150, model="gpt-4o-mini")
qa_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
    memory=memory
)

def query_documents(query: str, n_results: int = 3) -> List[Tuple[str, str]]:
    results = collection.query(query_texts=[query], n_results=n_results)
    return list(zip(results['documents'][0], results['metadatas'][0]))

def generate_response(query: str, chat_history: List[dict] = None) -> Tuple[str, List[dict]]:
    if chat_history is None:
        chat_history = []
    result = qa_chain({"question": query, "chat_history": chat_history})
    answer = result["answer"]
    relevant_docs = vectorstore.similarity_search(query, k=3)
    doc_references = []
    seen_files = set()

    for doc in relevant_docs:
        file_id = doc.metadata.get("file_id")
        if file_id and file_id not in seen_files:
            doc_references.append({
                "file_name": doc.metadata.get("file_name", "Unknown"),
                "source": doc.metadata.get("source", "Unknown path"),
                "file_id": file_id,
                "file_type": doc.metadata.get("file_type", "Unknown"),
                "relevance_score": doc.metadata.get("relevance_score", 0.0)
            })
            seen_files.add(file_id)
    return answer, doc_references

@app.route('/')
def index():
    if not session.get('token'):
        return redirect(url_for('login'))
    return render_template('onedrive.html', user=session.get('user', {}))

@app.route('/open_document/<file_id>')
def open_document(file_id):
    if not session.get('token'):
        return jsonify({"error": "Not authenticated"}), 401

    if not file_id or file_id == "Unknown":
        return jsonify({"error": "Invalid file ID"}), 400

    access_token = session['token']
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    try:
        # Attempt to get direct download URL
        response = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}', headers=headers)
        if response.status_code == 200:
            file_data = response.json()
            download_url = file_data.get('@microsoft.graph.downloadUrl')
            if download_url:
                return jsonify({"download_url": download_url})

        # Attempt to create sharing link
        sharing_body = {"type": "view", "scope": "organization"}
        sharing_response = requests.post(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/createLink', headers=headers, json=sharing_body)
        if sharing_response.status_code == 200:
            share_link = sharing_response.json().get('link', {}).get('webUrl')
            if share_link:
                return jsonify({"download_url": share_link})

        # Attempt to get a temporary access URL
        temp_url_response = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}?select=id,@microsoft.graph.downloadUrl', headers=headers)
        if temp_url_response.status_code == 200:
            temp_download_url = temp_url_response.json().get('@microsoft.graph.downloadUrl')
            if temp_download_url:
                return jsonify({"download_url": temp_download_url})

        # If all attempts fail, return detailed error
        return jsonify({
            "error": "Unable to generate access URL",
            "details": {
                "direct_url_status": response.status_code,
                "sharing_status": sharing_response.status_code,
                "file_id": file_id
            }
        }), 500

    except requests.exceptions.RequestException as e:
        return jsonify({"error": "Network error while accessing file"}), 500
    except Exception as e:
        return jsonify({"error": "Unexpected error while accessing file"}), 500

@app.route('/login')
def login():
    auth_url = msal_app.get_authorization_request_url(SCOPE, redirect_uri=REDIRECT_URI)
    return redirect(auth_url)

@app.route('/getAToken')
def authorized():
    if request.args.get('code'):
        token_result = msal_app.acquire_token_by_authorization_code(
            request.args['code'],
            scopes=SCOPE,
            redirect_uri=REDIRECT_URI
        )
        if "access_token" in token_result:
            session['user'] = token_result.get('id_token_claims')
            session['token'] = token_result['access_token']
            session.permanent = True
            try:
                processed, skipped = process_files_optimized(session['token'])
            except Exception as e:
                print(f"Error processing files: {str(e)}")
        else:
            print(f"Token acquisition failed: {token_result.get('error_description', 'Unknown error')}")
    return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
def upload_file():
    if not session.get('token'):
        return jsonify({'error': 'Authentication required'}), 403
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    try:
        filename = secure_filename(file.filename)
        file_stream = BytesIO()
        file.save(file_stream)
        file_stream.seek(0)
        file_size = file_stream.getbuffer().nbytes

        if file_size == 0:
            return jsonify({'error': 'Empty file'}), 400

        mime_type, _ = mimetypes.guess_type(filename)
        
        # Check if file is either a document or supported image
        is_doc = mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        is_image = ocr_processor.is_image_file(filename)
        
        if not (is_doc or is_image):
            return jsonify({'error': 'Only .docx files and images (jpg, png, tiff, bmp) are supported'}), 400

        headers = {'Authorization': f'Bearer {session["token"]}', 'Content-Type': 'application/json'}
        upload_session_url = f'https://graph.microsoft.com/v1.0/me/drive/root:/Documents/{filename}:/createUploadSession'
        upload_session_data = {"item": {"@microsoft.graph.conflictBehavior": "rename", "name": filename}}

        upload_session_response = requests.post(upload_session_url, headers=headers, json=upload_session_data)
        if not upload_session_response.ok:
            return jsonify({'error': 'Failed to create upload session', 'details': upload_session_response.text}), upload_session_response.status_code

        upload_url = upload_session_response.json().get('uploadUrl')
        if not upload_url:
            return jsonify({'error': 'No upload URL received'}), 500

        upload_headers = {'Content-Length': str(file_size), 'Content-Range': f'bytes 0-{file_size-1}/{file_size}'}
        file_stream.seek(0)
        upload_response = requests.put(upload_url, headers=upload_headers, data=file_stream)

        if not upload_response.ok:
            return jsonify({'error': 'Failed to upload file', 'details': upload_response.text}), upload_response.status_code

        file_id = upload_response.json().get('id')
        if not file_id:
            return jsonify({'error': 'No file ID received'}), 500

        # Process the file content based on type
        file_stream.seek(0)
        if is_doc:
            text_content = read_docx(file_stream)
        else:  # is_image
            text_content = ocr_processor.extract_text(file_stream.getvalue())

        if not text_content:
            return jsonify({'error': 'No content could be extracted from file'}), 400

        chunks = text_splitter.split_text(text_content)
        metadatas = [{
            "source": f"Documents/{filename}",
            "file_name": filename,
            "file_id": file_id,
            "file_type": mime_type,
            "chunk_index": i,
            "upload_date": datetime.now().isoformat()
        } for i in range(len(chunks))]
        
        vectorstore.add_texts(texts=chunks, metadatas=metadatas, ids=[f"{file_id}_chunk_{i}" for i in range(len(chunks))])

        return jsonify({
            'success': True,
            'message': 'File uploaded and processed successfully',
            'file_id': file_id,
            'file_name': filename,
            'file_type': mime_type,
            'chunks_processed': len(chunks)
        })

    except Exception as e:
        return jsonify({'error': 'Upload failed', 'details': str(e)}), 500

@app.route('/upload/status/<file_id>')
def check_upload_status(file_id):
    if not session.get('token'):
        return jsonify({'error': 'Not authenticated'}), 401

    headers = {'Authorization': f'Bearer {session["token"]}'}
    response = requests.get(f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}', headers=headers)

    if response.status_code == 200:
        return jsonify({'status': 'complete', 'file_info': response.json()})
    return jsonify({'status': 'pending'})

@app.route('/chat', methods=['POST'])
def chat():
    if not session.get('token'):
        return jsonify({"error": "Not authenticated"}), 401

    query = request.json.get('query')
    if not query:
        return jsonify({"error": "No query provided"}), 400

    chat_history = session.get('chat_history', [])
    result = qa_chain({"question": query, "chat_history": chat_history})
    response = result["answer"]
    relevant_docs = vectorstore.similarity_search(query, k=3)
    seen_file_ids = set()
    doc_references = []

    for doc in relevant_docs:
        file_id = doc.metadata.get("file_id")
        if file_id and file_id not in seen_file_ids:
            doc_references.append({
                "file_name": doc.metadata.get("file_name", "Unknown"),
                "source": doc.metadata.get("source", "Unknown path"),
                "file_id": file_id,
                "file_type": doc.metadata.get("file_type", "Unknown"),
                "relevance_score": doc.metadata.get("relevance_score", 0.0)
            })
            seen_file_ids.add(file_id)

    doc_references.sort(key=lambda x: x.get("relevance_score", 0), reverse=True)
    timestamp = datetime.now().isoformat()
    chat_history.append({"role": "user", "content": query, "timestamp": timestamp})
    chat_history.append({"role": "assistant", "content": response, "timestamp": timestamp, "documents": [doc["file_name"] for doc in doc_references]})
    session['chat_history'] = chat_history[-10:]

    return jsonify({"response": response, "relevant_documents": doc_references, "timestamp": timestamp})

@app.route('/delete_file/<file_id>', methods=['DELETE'])
def delete_file(file_id):
    if not session.get('token'):
        return jsonify({"error": "Not authenticated"}), 401

    headers = {'Authorization': f'Bearer {session["token"]}'}
    
    try:
        # First verify the file exists
        verify_response = requests.get(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}',
            headers=headers
        )
        
        if verify_response.status_code == 404:
            return jsonify({
                "success": False,
                "error": "File not found or already deleted"
            }), 404
        
        if not verify_response.ok:
            return jsonify({
                "success": False,
                "error": f"Failed to verify file: {verify_response.json().get('error', {}).get('message', 'Unknown error')}"
            }), verify_response.status_code

        # Proceed with deletion if file exists
        delete_response = requests.delete(
            f'https://graph.microsoft.com/v1.0/me/drive/items/{file_id}',
            headers=headers
        )
        
        if delete_response.status_code == 204:  # Successful deletion
            try:
                # Remove vectors from vectorstore
                vectorstore.delete(filter={"file_id": file_id})
            except Exception as ve:
                print(f"Warning: Failed to delete vectors: {str(ve)}")
            
            return jsonify({
                "success": True,
                "message": "File deleted successfully"
            })
        else:
            error_message = "Unknown error"
            if delete_response.content:
                try:
                    error_data = delete_response.json()
                    error_message = error_data.get('error', {}).get('message', 'Unknown error')
                except:
                    pass
                    
            return jsonify({
                "success": False,
                "error": f"OneDrive API error: {error_message}",
                "status": delete_response.status_code
            }), delete_response.status_code
            
    except requests.exceptions.RequestException as e:
        return jsonify({
            "success": False,
            "error": "Failed to connect to OneDrive API",
            "details": str(e)
        }), 500
    except Exception as e:
        return jsonify({
            "success": False,
            "error": "Server error while deleting file",
            "details": str(e)
        }), 500

@app.route('/check_auth')
def check_auth():
    if 'token' in session and 'user' in session:
        return jsonify({'authenticated': True})
    return jsonify({'authenticated': False}), 401

@app.errorhandler(401)
def unauthorized(e):
    session.clear()
    return redirect(url_for('login'))

if __name__ == "__main__":
    app.run(debug=True)  # You may want to set debug=False for production
